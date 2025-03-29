import os
import random
import re
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx import Document
import logging
import glob

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    filename='invoice_automation.log'
)
logger = logging.getLogger('invoice_automation')

def ordinal(n):
    """Return the ordinal representation of a number (1st, 2nd, 3rd, etc.)"""
    if 10 <= n % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
    return f"{n}{suffix}"

def get_month_name(month_num):
    """Convert month number to full month name"""
    return datetime.strptime(str(month_num), "%m").strftime("%B")

def get_financial_year(date):
    """Get financial year in format YYYY-YY based on date"""
    if date.month >= 4:  # April onwards is new financial year in India
        return f"{date.year}-{str(date.year + 1)[-2:]}"
    else:
        return f"{date.year - 1}-{str(date.year)[-2:]}"

def find_latest_invoice():
    """Find the latest invoice file based on naming pattern"""
    invoice_files = glob.glob("invoice_*.docx")
    if not invoice_files:
        return None
    
    # Sort files by modification time (newest first)
    latest_file = max(invoice_files, key=os.path.getmtime)
    logger.info(f"Found latest invoice file: {latest_file}")
    return latest_file

def extract_date_from_latest_invoice(latest_file):
    """Extract the invoice date from the latest invoice file"""
    try:
        doc = Document(latest_file)
        for para in doc.paragraphs:
            if "Date:" in para.text and ("st " in para.text or "nd " in para.text or "rd " in para.text or "th " in para.text):
                # Extract date like "1st February 2025"
                pattern = r'\d{1,2}(st|nd|rd|th)\s+([A-Za-z]+)\s+(\d{4})'
                match = re.search(pattern, para.text)
                if match:
                    month_name = match.group(2)
                    year = int(match.group(3))
                    # Create a date object for the first day of that month
                    date_str = f"1 {month_name} {year}"
                    try:
                        date = datetime.strptime(date_str, "%d %B %Y")
                        logger.info(f"Extracted date from latest invoice: {date}")
                        return date
                    except ValueError:
                        logger.error(f"Could not parse date: {date_str}")
        
        logger.warning("Could not extract date from latest invoice")
        return None
    except Exception as e:
        logger.error(f"Error extracting date from latest invoice: {str(e)}")
        return None

def determine_invoice_date():
    """Determine the invoice date for the new invoice"""
    # First try to find the latest invoice
    latest_file = find_latest_invoice()
    
    if latest_file:
        # Extract date from the latest invoice
        last_date = extract_date_from_latest_invoice(latest_file)
        if last_date:
            # Set new invoice date to the next month
            next_date = last_date + relativedelta(months=1)
            logger.info(f"Setting new invoice date to next month: {next_date}")
            return next_date
    
    # If no previous invoice found or date couldn't be extracted, use current date
    current_date = datetime.now().replace(day=1)
    logger.info(f"Using current month for invoice date: {current_date}")
    return current_date

def auto_update_invoice(template_name):
    """Automatically update invoice with next month's information"""
    try:
        # Load the document
        doc = Document(template_name)
        
        # Determine invoice date (first day of the month)
        invoice_date = determine_invoice_date()
        
        # Calculate the financial year
        fin_year = get_financial_year(invoice_date)
        
        # Generate random invoice number
        random_num = random.randint(1000, 9999)
        invoice_number = f"IN{random_num}/{fin_year}"
        
        # Calculate service dates (7th of previous month to 6th of current month)
        prev_month_date = invoice_date - relativedelta(months=1)
        service_start_date = prev_month_date.replace(day=7)
        service_end_date = invoice_date.replace(day=6)
        
        # Format dates
        date_format1 = f"{ordinal(invoice_date.day)} {invoice_date.strftime('%B %Y')}"  # 1st February 2025
        date_format2 = f"{invoice_date.day:d}/{invoice_date.month:02d}/{invoice_date.year}"  # 1/02/2025
        service_period = f"{ordinal(service_start_date.day)} {service_start_date.strftime('%B %Y')} to {ordinal(service_end_date.day)} {service_end_date.strftime('%B %Y')}"
        
        logger.info(f"Updating document with: Invoice #{invoice_number}, Date: {date_format1}, Service Period: {service_period}")
        
        # Update text in the document
        date_format1_updated = False
        date_format2_updated = False
        invoice_number_updated = False
        service_period_updated = False
        
        # Handle the main invoice date (1st February 2025 format)
        for para in doc.paragraphs:
            if "Date:" in para.text and ("st" in para.text or "nd" in para.text or "rd" in para.text or "th" in para.text) and not date_format1_updated:
                # Look for text like "Date: 1st February 2025"
                pattern = r'Date:\s+\d{1,2}(st|nd|rd|th)\s+[A-Za-z]+\s+\d{4}'
                match = re.search(pattern, para.text)
                if match:
                    original_text = para.text
                    para.text = re.sub(pattern, f"Date: {date_format1}", para.text)
                    if original_text != para.text:
                        date_format1_updated = True
                        logger.info("Updated date format 1 (1st Month Year)")
        
        # Handle the tables - both for the invoice number, date and service period
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Look for Invoice Number
                        if "Invoice No:" in paragraph.text and not invoice_number_updated:
                            pattern = r'IN\d{4}/\d{4}-\d{2}'
                            match = re.search(pattern, paragraph.text)
                            if match:
                                original_text = paragraph.text
                                paragraph.text = re.sub(pattern, invoice_number, paragraph.text)
                                if original_text != paragraph.text:
                                    invoice_number_updated = True
                                    logger.info("Updated invoice number")
                        
                        # Look for date in format "Date: - 1/02/2025"
                        if "Date:" in paragraph.text and "/" in paragraph.text and not date_format2_updated:
                            pattern = r'\d{1,2}/\d{2}/\d{4}'
                            match = re.search(pattern, paragraph.text)
                            if match:
                                original_text = paragraph.text
                                paragraph.text = re.sub(pattern, date_format2, paragraph.text)
                                if original_text != paragraph.text:
                                    date_format2_updated = True
                                    logger.info("Updated date format 2 (DD/MM/YYYY)")
                        
                        # Look for service period
                        pattern = r'\d{1,2}(st|nd|rd|th)\s+[A-Za-z]+\s+\d{4}\s+to\s+\d{1,2}(st|nd|rd|th)\s+[A-Za-z]+\s+\d{4}'
                        match = re.search(pattern, paragraph.text)
                        if match and not service_period_updated:
                            original_text = paragraph.text
                            paragraph.text = re.sub(pattern, service_period, paragraph.text)
                            if original_text != paragraph.text:
                                service_period_updated = True
                                logger.info("Updated service period")
        
        # Additional check for invoice number and date in paragraphs
        if not invoice_number_updated or not date_format2_updated:
            for para in doc.paragraphs:
                # Look for invoice number
                if not invoice_number_updated and "Invoice No:" in para.text:
                    pattern = r'IN\d{4}/\d{4}-\d{2}'
                    match = re.search(pattern, para.text)
                    if match:
                        original_text = para.text
                        para.text = re.sub(pattern, invoice_number, para.text)
                        if original_text != para.text:
                            invoice_number_updated = True
                            logger.info("Updated invoice number in paragraph")
                
                # Look for date in DD/MM/YYYY format
                if not date_format2_updated and "Date:" in para.text and "/" in para.text:
                    pattern = r'\d{1,2}/\d{2}/\d{4}'
                    match = re.search(pattern, para.text)
                    if match:
                        original_text = para.text
                        para.text = re.sub(pattern, date_format2, para.text)
                        if original_text != para.text:
                            date_format2_updated = True
                            logger.info("Updated date format 2 in paragraph")
        
        # Log warnings for any updates that weren't made
        if not date_format1_updated:
            logger.warning("Could not find or update date format 1 (1st Month Year)")
        if not date_format2_updated:
            logger.warning("Could not find or update date format 2 (DD/MM/YYYY)")
        if not invoice_number_updated:
            logger.warning("Could not find or update invoice number")
        if not service_period_updated:
            logger.warning("Could not find or update service period")
        
        # Create output filename based on month name and year
        output_filename = f"invoice_{invoice_date.strftime('%B_%Y').lower()}.docx"
        
        # Save the document
        doc.save(output_filename)
        logger.info(f"Invoice updated successfully and saved as {output_filename}")
        return output_filename
        
    except Exception as e:
        logger.error(f"Error updating invoice: {str(e)}")
        return None

if __name__ == "__main__":
    # This could be triggered by a scheduled job in the cloud
    logger.info("Starting automated invoice generation")
    
    # Look for template file
    template_files = glob.glob("*template*.docx")
    if not template_files:
        # Fall back to the first .docx file we find if there's no obvious template
        template_files = glob.glob("*.docx")
        if not template_files:
            logger.error("No template file found. Exiting.")
            exit(1)
    
    template_name = template_files[0]
    logger.info(f"Using template file: {template_name}")
    
    result = auto_update_invoice(template_name)
    if result:
        logger.info(f"Successfully generated invoice: {result}")
    else:
        logger.error("Failed to generate invoice")
